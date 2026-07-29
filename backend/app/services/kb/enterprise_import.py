"""企业规范导入工具（SubTask 14.1）。

支持 PDF / Word(.docx) / Excel(.xlsx) 三种格式，输出统一的结构化条文
``ClauseRecord`` 列表，与已有 KB schema 一致，便于直接写入 Qdrant 索引。

遵循"八荣八耻"原则：
- 以复用现有为荣：PDF 提取复用 ``kb/tools/extract_clauses.py`` 中的
  pdfplumber 调用模式与正则切分逻辑（``_CLAUSE_PATTERN``、
  ``_guess_category``、``_auto_keywords``、``_extract_references``）。
- 以瞎猜接口为耻：所有第三方库（pdfplumber / python-docx / openpyxl）
  均先验证可用性再使用，缺失时抛 ``ImportError`` 并提示安装命令。
- 以最小修改为荣：不修改已有 indexer/retriever，仅新增本模块。
- 降级路径：不支持的格式抛 ``ValueError`` 并提示支持的格式列表。
"""

from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Any

from app.logging import get_logger
from app.schemas.kb import ClauseRecord

log = get_logger(__name__)


# ---------------------------------------------------------------------------
# 复用 kb/tools/extract_clauses.py 中的正则与辅助函数
# （行内复制以避免 sys.path 注入到 tools 目录；与原文件保持一致语义）
# ---------------------------------------------------------------------------

# 匹配条款标题：行首数字编号 + 空格 + 标题文本，如 "5.2 圆度公差" / "4.1.1 尺寸界线"
_CLAUSE_PATTERN = re.compile(
    r"^(?P<id>\d+(?:\.\d+){0,3})\s+(?P<title>[^\n]{2,80})\s*$",
    re.MULTILINE,
)

_REF_PATTERN = re.compile(
    r"(GB/T\s*\d+(?:\.\d+)?(?:[-—]\d{4})?|ISO\s*\d+(?:[-—]\d{4})?|QB/\w+\s*\d+(?:[-—]\d{4})?|JB/T\s*\d+(?:[-—]\d{4})?)"
)

_CATEGORY_KEYWORDS = {
    "shape_tolerance": ["圆度", "圆柱度", "平面度", "直线度", "形状公差"],
    "orientation_tolerance": ["平行度", "垂直度", "倾斜度", "方向公差"],
    "location_tolerance": ["位置度", "同轴度", "对称度", "位置公差"],
    "runout_tolerance": ["圆跳动", "全跳动", "跳动公差"],
    "dimension_basic": ["尺寸", "标注", "基本规则"],
    "dimension_line": ["尺寸界线", "尺寸线", "箭头"],
    "surface_parameter": ["表面结构", "Ra", "Rz", "粗糙度"],
    "surface_symbol": ["图形符号", "表面结构符号"],
    "line_type": ["实线", "虚线", "点画线", "图线"],
    "general_tolerance": ["一般公差", "未注公差"],
    "cad_layer": ["图层", "CAD"],
    "cad_dimension": ["CAD", "尺寸标注"],
    "enterprise_specific": ["企业", "本公司", "厂内", "内控"],
}


def _guess_category(title: str, body: str) -> str:
    text = (title + " " + body).lower()
    for cat, kws in _CATEGORY_KEYWORDS.items():
        if any(kw in text for kw in kws):
            return cat
    return "general"


def _auto_keywords(title: str) -> list[str]:
    parts = re.split(r"[\s、，,。/（）()]+", title)
    return [p for p in parts if len(p) >= 2 and not p.replace(".", "").isdigit()][:5]


def _extract_version(standard: str) -> str:
    m = re.search(r"(\d{4})$", standard)
    return m.group(1) if m else ""


def _extract_references(body: str, self_standard: str) -> list[str]:
    refs: list[str] = []
    for m in _REF_PATTERN.finditer(body):
        ref = m.group(1).replace("—", "-").replace("  ", " ")
        ref = re.sub(r"\s+", " ", ref).strip()
        if ref and ref not in refs and ref not in self_standard:
            refs.append(ref)
    return refs[:5]


def _split_into_clauses(
    full_text: str, standard: str, version: str, source_file: str
) -> list[ClauseRecord]:
    """按正则识别条款标题，切分正文为 ClauseRecord 列表。"""
    matches = list(_CLAUSE_PATTERN.finditer(full_text))
    if not matches:
        # 未识别到条款，整体作为单条条款返回，clause_id 置 "0"
        body = full_text.strip()[:2000]
        if not body:
            return []
        return [
            ClauseRecord(
                standard=standard,
                clause_id="0",
                title=f"{standard} 提取内容",
                category="general",
                keywords=_auto_keywords(standard),
                references=[],
                version=version,
                is_sample=False,
                original_text=body,
                source_file=source_file,
            )
        ]

    records: list[ClauseRecord] = []
    for i, m in enumerate(matches):
        clause_id = m.group("id")
        title = m.group("title").strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
        body = full_text[start:end].strip()
        if not body:
            continue
        records.append(
            ClauseRecord(
                standard=standard,
                clause_id=clause_id,
                title=title,
                category=_guess_category(title, body),
                keywords=_auto_keywords(title),
                references=_extract_references(body, standard),
                version=version,
                is_sample=False,
                original_text=body,
                source_file=source_file,
            )
        )
    return records


# ---------------------------------------------------------------------------
# 格式分发
# ---------------------------------------------------------------------------


SUPPORTED_FORMATS = ("pdf", "docx", "xlsx")
"""支持的文件格式（小写，不含点）。"""


def _detect_format(file_path: Path) -> str:
    """根据扩展名识别格式。不支持时抛 ValueError。"""
    ext = file_path.suffix.lower().lstrip(".")
    if ext in SUPPORTED_FORMATS:
        return ext
    raise ValueError(
        f"不支持的企业规范文件格式：{file_path.suffix!r}。"
        f"支持的格式：{', '.join(SUPPORTED_FORMATS)}"
    )


def _check_pdfplumber() -> None:
    try:
        import pdfplumber  # noqa: F401  # noqa: F811
    except ImportError as e:  # pragma: no cover
        raise ImportError(
            "pdfplumber 未安装，请运行：pip install pdfplumber==0.11.10"
        ) from e


def _check_docx() -> None:
    try:
        import docx  # type: ignore[import-not-found]  # noqa: F401
    except ImportError as e:  # pragma: no cover
        raise ImportError(
            "python-docx 未安装，请运行：pip install python-docx==1.2.0"
        ) from e


def _check_openpyxl() -> None:
    try:
        import openpyxl  # type: ignore[import-not-found]  # noqa: F401
    except ImportError as e:  # pragma: no cover
        raise ImportError(
            "openpyxl 未安装，请运行：pip install openpyxl==3.1.5"
        ) from e


# ---------------------------------------------------------------------------
# 各格式提取器
# ---------------------------------------------------------------------------


def _extract_from_pdf(
    pdf_path: Path, standard: str, version: str
) -> list[ClauseRecord]:
    """从 PDF 提取条款（复用 kb/tools/extract_clauses.py 的 pdfplumber 模式）。"""
    _check_pdfplumber()
    import pdfplumber  # type: ignore[import-not-found]

    full_text_parts: list[str] = []
    all_tables: list[list[list[str]]] = []

    with pdfplumber.open(str(pdf_path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            full_text_parts.append(page_text)
            try:
                page_tables = page.extract_tables() or []
            except Exception:  # noqa: BLE001
                page_tables = []
            for tbl in page_tables:
                if tbl and len(tbl) >= 1:
                    all_tables.append(
                        [[(cell or "").strip() for cell in row] for row in tbl]
                    )

    full_text = "\n".join(full_text_parts)

    # 表格附加到全文末尾，便于条款切分时纳入 references
    if all_tables:
        tables_md_parts: list[str] = ["\n\n## 表格汇总\n"]
        for idx, table in enumerate(all_tables, start=1):
            tables_md_parts.append(f"\n### 表 {idx}\n")
            tables_md_parts.append(_table_to_markdown(table))
        full_text += "\n".join(tables_md_parts)

    records = _split_into_clauses(
        full_text, standard=standard, version=version, source_file=pdf_path.name
    )
    log.info(
        "kb.enterprise_import.pdf_extracted",
        file=str(pdf_path),
        standard=standard,
        clauses=len(records),
        tables=len(all_tables),
    )
    return records


def _extract_from_docx(
    docx_path: Path, standard: str, version: str
) -> list[ClauseRecord]:
    """从 Word(.docx) 提取条款。

    使用 python-docx 读取段落与表格，合并为纯文本后走通用切分逻辑。
    """
    _check_docx()
    import docx  # type: ignore[import-not-found]

    doc = docx.Document(str(docx_path))
    parts: list[str] = []
    table_markdown_parts: list[str] = []

    # 段落
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)

    # 表格 → Markdown
    for tbl_idx, table in enumerate(doc.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            rows.append(cells)
        if rows:
            table_markdown_parts.append(f"\n### 表 {tbl_idx}\n")
            table_markdown_parts.append(_table_to_markdown(rows))

    full_text = "\n".join(parts)
    if table_markdown_parts:
        full_text += "\n\n## 表格汇总\n" + "\n".join(table_markdown_parts)

    records = _split_into_clauses(
        full_text, standard=standard, version=version, source_file=docx_path.name
    )
    log.info(
        "kb.enterprise_import.docx_extracted",
        file=str(docx_path),
        standard=standard,
        clauses=len(records),
        tables=len(doc.tables),
    )
    return records


def _extract_from_xlsx(
    xlsx_path: Path, standard: str, version: str
) -> list[ClauseRecord]:
    """从 Excel(.xlsx) 提取条款。

    约定：每行一条条款，列布局为
    [条款号, 标题, 正文, 关键词(逗号分隔), 引用(逗号分隔)]。
    缺失列容忍；首行若为表头（含"条款号"等中文字样）自动跳过。
    """
    _check_openpyxl()
    import openpyxl  # type: ignore[import-not-found]

    wb = openpyxl.load_workbook(str(xlsx_path), data_only=True)
    records: list[ClauseRecord] = []

    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        # 检测首行是否为表头
        first_row = [str(c).strip() if c is not None else "" for c in rows[0]]
        first_row_text = "|".join(first_row).lower()
        if any(
            kw in first_row_text
            for kw in ["条款号", "条款编号", "clause", "标题", "title"]
        ):
            data_rows = rows[1:]
        else:
            data_rows = rows

        for row in data_rows:
            cells = [
                (str(c).strip() if c is not None else "") for c in row
            ]
            if not any(cells):
                continue
            # 至少需要条款号或标题或正文
            if len(cells) < 3:
                # 行数据不足，将整行拼成正文做兜底
                clause_id = ""
                title = standard
                body = " | ".join(cells)
                keywords: list[str] = []
                references: list[str] = []
            else:
                clause_id = cells[0] or ""
                title = cells[1] or cells[0] or standard
                body = cells[2] or ""
                keywords = [
                    k.strip()
                    for k in (cells[3].split(",") if len(cells) > 3 and cells[3] else [])
                    if k.strip()
                ]
                references = [
                    r.strip()
                    for r in (cells[4].split(",") if len(cells) > 4 and cells[4] else [])
                    if r.strip()
                ]

            if not body.strip():
                continue

            if not keywords:
                keywords = _auto_keywords(title)
            if not references:
                references = _extract_references(body, standard)

            records.append(
                ClauseRecord(
                    standard=standard,
                    clause_id=clause_id or f"row-{len(records) + 1}",
                    title=title,
                    category=_guess_category(title, body),
                    keywords=keywords,
                    references=references,
                    version=version,
                    is_sample=False,
                    original_text=body,
                    source_file=xlsx_path.name,
                )
            )

    log.info(
        "kb.enterprise_import.xlsx_extracted",
        file=str(xlsx_path),
        standard=standard,
        clauses=len(records),
        sheets=len(wb.worksheets),
    )
    return records


def _table_to_markdown(table: list[list[str]]) -> str:
    if not table:
        return ""
    header = table[0]
    rows = table[1:] if len(table) > 1 else []
    md = "| " + " | ".join(header) + " |\n"
    md += "|" + "---|" * len(header) + "\n"
    for row in rows:
        row = row + [""] * (len(header) - len(row))
        md += "| " + " | ".join(row[: len(header)]) + " |\n"
    return md


# ---------------------------------------------------------------------------
# 公共入口
# ---------------------------------------------------------------------------


def import_enterprise_standard(
    file_path: str | Path,
    standard_name: str,
    version: str = "",
) -> list[ClauseRecord]:
    """导入企业规范文件，返回结构化条款列表。

    Args:
        file_path: 文件路径（PDF / Word / Excel）
        standard_name: 规范编号或名称，如 "Q/XX 001-2024" 或 "企业内控规范"
        version: 规范版本年份，留空则从 standard_name 末尾年份推断

    Returns:
        list[ClauseRecord]：与已有 KB schema 一致的结构化条款

    Raises:
        ValueError: 文件格式不支持
        FileNotFoundError: 文件不存在
        ImportError: 所需第三方库未安装
    """
    p = Path(file_path)
    if not p.is_file():
        raise FileNotFoundError(f"企业规范文件不存在：{p}")

    fmt = _detect_format(p)
    if not version:
        version = _extract_version(standard_name)

    log.info(
        "kb.enterprise_import.start",
        file=str(p),
        format=fmt,
        standard=standard_name,
        version=version,
    )

    if fmt == "pdf":
        records = _extract_from_pdf(p, standard_name, version)
    elif fmt == "docx":
        records = _extract_from_docx(p, standard_name, version)
    elif fmt == "xlsx":
        records = _extract_from_xlsx(p, standard_name, version)
    else:  # pragma: no cover  # _detect_format 已校验
        raise ValueError(f"未实现的格式：{fmt}")

    log.info(
        "kb.enterprise_import.done",
        file=str(p),
        standard=standard_name,
        clauses=len(records),
    )
    return records


def main(argv: list[str] | None = None) -> int:
    """命令行入口：python -m app.services.kb.enterprise_import <file> <standard> [version]"""
    args = argv if argv is not None else sys.argv[1:]
    if len(args) < 2:
        sys.stderr.write(
            "用法: python -m app.services.kb.enterprise_import "
            "<file_path> <standard_name> [version]\n"
        )
        return 2
    file_path = args[0]
    standard = args[1]
    version = args[2] if len(args) > 2 else ""
    try:
        records = import_enterprise_standard(file_path, standard, version)
    except (ValueError, FileNotFoundError, ImportError) as e:
        sys.stderr.write(f"[enterprise_import] 失败：{e}\n")
        return 1
    print(f"已提取 {len(records)} 条条款")
    for r in records[:5]:
        print(f"  - {r.standard} §{r.clause_id} {r.title}")
    if len(records) > 5:
        print(f"  ... 共 {len(records)} 条")
    return 0


if __name__ == "__main__":  # pragma: no cover
    sys.exit(main())
