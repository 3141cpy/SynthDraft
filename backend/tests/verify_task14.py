"""Task 14 企业规范自定义 端到端实测脚本。

覆盖 SubTask：
- 14.1: 企业规范导入工具（PDF / Word / Excel 各 1 例 + 不支持格式异常 1 例）
- 14.2: 规范冲突检测（构造 2 条矛盾条文，验证检测到 conflict_type=contradiction）
- 14.3: 多套规范配置切换（创建 2 个 profile，切换 active，验证检索使用新 profile）
- 降级路径：LLM 不可用时关键词匹配仍工作；PostgreSQL 不可用时降级 JSON 文件

运行：
    cd d:\\SynthDraft\\backend
    .venv\\Scripts\\python.exe tests\\verify_task14.py

设计原则（八荣八耻）：
- 复用现有：通过 sys.path 注入 backend 目录，复用现有 settings 与 KB 模块
- 实事求是：环境限制项（无 Qdrant / 无 LLM 服务）如实标注，仅验证降级路径
- 覆盖测试：每个 SubTask 必须有可执行的断言
"""

from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path
from typing import Any

# 将 backend 目录加入 sys.path
BACKEND_ROOT = Path(__file__).resolve().parents[1]
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

# 测试环境变量（避免污染真实配置）
os.environ.setdefault("LOG_LEVEL", "WARNING")
os.environ.setdefault("OFFLINE_MODE", "false")


def section(title: str) -> None:
    print(f"\n{'=' * 70}\n{title}\n{'-' * 70}", flush=True)


def _ok(msg: str) -> None:
    print(f"  [PASS] {msg}", flush=True)


def _fail(msg: str, detail: str = "") -> None:
    print(f"  [FAIL] {msg}{f' :: {detail}' if detail else ''}", flush=True)


def _info(msg: str) -> None:
    print(f"  [INFO] {msg}", flush=True)


def _env_limit(msg: str) -> None:
    print(f"  [ENV-LIMIT] {msg}", flush=True)


# ===== 全局统计 =====
_passed = 0
_failed = 0
_env_limits = 0
_failures: list[str] = []


def check(condition: bool, msg: str, detail: str = "") -> None:
    global _passed, _failed
    if condition:
        _passed += 1
        _ok(msg)
    else:
        _failed += 1
        _failures.append(f"{msg}{f' :: {detail}' if detail else ''}")
        _fail(msg, detail)


def env_limit(msg: str) -> None:
    global _env_limits
    _env_limits += 1
    _env_limit(msg)


# ---------------------------------------------------------------------------
# 测试夹具：构造测试文件
# ---------------------------------------------------------------------------


def _make_test_pdf(pdf_path: Path) -> None:
    """用 reportlab 生成一个最小 PDF 测试文件（仅含文本）。

    若 reportlab 不可用，则用 pypdf 写入一个简单 PDF。
    """
    try:
        from reportlab.pdfgen import canvas  # type: ignore[import-not-found]
        from reportlab.pdfbase import pdfmetrics  # type: ignore[import-not-found]
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont  # type: ignore[import-not-found]

        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        c = canvas.Canvas(str(pdf_path))
        c.setFont("STSong-Light", 14)
        c.drawString(80, 750, "5.1 直线度公差")
        c.drawString(80, 730, "直线度公差带是距离为公差值 0.05 的两平行直线之间的区域。")
        c.drawString(80, 700, "5.2 圆度公差")
        c.drawString(80, 680, "圆度公差带是在同一正截面上，半径差为公差值 0.1 的两同心圆之间的区域。")
        c.showPage()
        c.save()
        return
    except ImportError:
        pass
    # reportlab 不可用 → 用 pdfplumber 间接生成（无字体支持，仅占位）
    # 兜底：用 pypdf 写入纯 ASCII 文本
    try:
        from pypdf import PdfWriter  # type: ignore[import-not-found]

        writer = PdfWriter()
        # pypdf 不能直接写文本，需空白页；这里仅生成空 PDF 作为格式占位
        writer.add_blank_page(width=595, height=842)
        with open(pdf_path, "wb") as f:
            writer.write(f)
        _info(f"用 pypdf 生成空 PDF（无文本）：{pdf_path}")
    except Exception as e:  # noqa: BLE001
        _info(f"生成 PDF 失败：{e}，测试将仅验证文件可被打开")


def _make_test_docx(docx_path: Path) -> None:
    """生成一个最小 Word(.docx) 测试文件。"""
    import docx  # type: ignore[import-not-found]

    doc = docx.Document()
    doc.add_heading("企业内控规范", level=1)
    doc.add_paragraph("5.1 直线度公差")
    doc.add_paragraph("直线度公差带是距离为公差值 0.05 的两平行直线之间的区域。")
    doc.add_paragraph("5.2 圆度公差")
    doc.add_paragraph("圆度公差带是在同一正截面上，半径差为公差值 0.1 的两同心圆之间的区域。")
    doc.add_paragraph("6.1 表面粗糙度")
    doc.add_paragraph("企业标准要求 Ra 不得大于 1.6，严于国标建议值。")
    doc.save(str(docx_path))


def _make_test_xlsx(xlsx_path: Path) -> None:
    """生成一个最小 Excel(.xlsx) 测试文件。"""
    import openpyxl  # type: ignore[import-not-found]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "条款"
    ws.append(["条款号", "标题", "正文", "关键词", "引用"])
    ws.append(["5.1", "直线度公差", "直线度公差带是距离为公差值 0.05 的两平行直线之间的区域。", "直线度,公差", "GB/T 1182-2018"])
    ws.append(["5.2", "圆度公差", "圆度公差带是在同一正截面上，半径差为公差值 0.1 的两同心圆之间的区域。", "圆度,公差", "GB/T 1182-2018"])
    ws.append(["6.1", "表面粗糙度", "Ra 不得大于 1.6，必须严格执行。", "Ra,粗糙度", "GB/T 131-2006"])
    wb.save(str(xlsx_path))


# ---------------------------------------------------------------------------
# SubTask 14.1：企业规范导入
# ---------------------------------------------------------------------------


def test_enterprise_import() -> None:
    section("SubTask 14.1: 企业规范导入工具")

    # 1. 模块导入
    try:
        from app.services.kb.enterprise_import import (
            import_enterprise_standard,
            SUPPORTED_FORMATS,
            _detect_format,
        )
        check(True, "enterprise_import 模块导入成功")
    except Exception as e:
        check(False, "enterprise_import 模块导入失败", str(e))
        return

    # 2. 支持格式列表
    check(
        set(SUPPORTED_FORMATS) == {"pdf", "docx", "xlsx"},
        f"SUPPORTED_FORMATS 含 pdf/docx/xlsx（实际 {SUPPORTED_FORMATS}）",
    )

    # 3. 不支持格式异常
    fake_file = Path("not_supported.txt")
    try:
        _detect_format(fake_file)
        check(False, "不支持格式应抛 ValueError")
    except ValueError as e:
        check(True, "不支持格式抛 ValueError", str(e)[:80])

    # 4. 文件不存在异常
    try:
        import_enterprise_standard("nonexistent.pdf", "Q/XX 001-2024")
        check(False, "文件不存在应抛 FileNotFoundError")
    except FileNotFoundError as e:
        check(True, "文件不存在抛 FileNotFoundError", str(e)[:80])

    # 5. Word 导入（最可靠，不依赖 PDF 字体）
    try:
        import docx  # type: ignore[import-not-found]  # noqa: F401
        with tempfile.TemporaryDirectory(prefix="task14_docx_") as tmpdir:
            docx_path = Path(tmpdir) / "test.docx"
            _make_test_docx(docx_path)
            records = import_enterprise_standard(
                docx_path, "Q/ENT 001-2024", version="2024"
            )
            check(
                len(records) >= 2,
                f"Word 导入提取 >= 2 条条款（实际 {len(records)}）",
                f"clauses={[r.title for r in records[:5]]}",
            )
            if records:
                check(
                    all(r.standard == "Q/ENT 001-2024" for r in records),
                    "所有条款 standard 字段正确",
                )
                check(
                    all(r.version == "2024" for r in records),
                    "所有条款 version 字段正确",
                )
                check(
                    all(r.source_file == "test.docx" for r in records),
                    "所有条款 source_file 正确",
                )
                check(
                    all(not r.is_sample for r in records),
                    "企业规范 is_sample=False",
                )
                check(
                    any(r.clause_id for r in records),
                    "至少有一条条款有 clause_id",
                )
                # 含表格汇总也应被切分
                check(
                    any("直线度" in r.title or "圆度" in r.title for r in records),
                    "提取出直线度/圆度条款",
                )
    except ImportError as e:
        env_limit(f"python-docx 未安装：{e}")
    except Exception as e:
        check(False, "Word 导入测试异常", str(e))

    # 6. Excel 导入
    try:
        import openpyxl  # type: ignore[import-not-found]  # noqa: F401
        with tempfile.TemporaryDirectory(prefix="task14_xlsx_") as tmpdir:
            xlsx_path = Path(tmpdir) / "test.xlsx"
            _make_test_xlsx(xlsx_path)
            records = import_enterprise_standard(
                xlsx_path, "Q/ENT 002-2024", version="2024"
            )
            check(
                len(records) == 3,
                f"Excel 导入提取 3 条条款（实际 {len(records)}）",
                f"clauses={[r.clause_id for r in records]}",
            )
            if records:
                # 验证表头跳过
                check(
                    all(r.title != "标题" for r in records),
                    "首行表头被正确跳过",
                )
                check(
                    records[0].clause_id == "5.1",
                    f"第一条 clause_id=5.1（实际 {records[0].clause_id}）",
                )
                check(
                    "直线度" in records[0].keywords or "公差" in records[0].keywords,
                    f"第一条关键词包含 直线度/公差（实际 {records[0].keywords}）",
                )
                check(
                    "GB/T 1182-2018" in records[0].references,
                    f"第一条引用 GB/T 1182-2018（实际 {records[0].references}）",
                )
                # 含"必须"的条文应归类 enterprise_specific 或 general
                last = records[-1]
                check(
                    "必须" in last.original_text,
                    "最后一条正文含 '必须' 措辞",
                )
    except ImportError as e:
        env_limit(f"openpyxl 未安装：{e}")
    except Exception as e:
        check(False, "Excel 导入测试异常", str(e))

    # 7. PDF 导入（可能受字体限制，环境限制如实标注）
    try:
        import pdfplumber  # type: ignore[import-not-found]  # noqa: F401
        with tempfile.TemporaryDirectory(prefix="task14_pdf_") as tmpdir:
            pdf_path = Path(tmpdir) / "test.pdf"
            _make_test_pdf(pdf_path)
            records = import_enterprise_standard(
                pdf_path, "Q/ENT 003-2024", version="2024"
            )
            # PDF 文本提取依赖字体支持，可能为 0 条
            if len(records) == 0:
                env_limit("PDF 文本提取为 0 条（可能 reportlab 未装，仅生成空 PDF）")
            else:
                check(
                    len(records) >= 1,
                    f"PDF 导入提取 >= 1 条条款（实际 {len(records)}）",
                )
                check(
                    all(r.source_file == "test.pdf" for r in records),
                    "所有条款 source_file=test.pdf",
                )
    except ImportError as e:
        env_limit(f"pdfplumber 未安装：{e}")
    except Exception as e:
        check(False, "PDF 导入测试异常", str(e))


# ---------------------------------------------------------------------------
# SubTask 14.2：规范冲突检测
# ---------------------------------------------------------------------------


def test_conflict_detection() -> None:
    section("SubTask 14.2: 规范冲突检测")

    # 1. 模块导入
    try:
        from app.services.kb.conflict_detector import (
            detect_conflicts,
            _keyword_detect_pair,
            _has_number_contradiction,
            _has_strict_vs_weaker,
            _jaccard,
            _tokenize_zh,
        )
        check(True, "conflict_detector 模块导入成功")
    except Exception as e:
        check(False, "conflict_detector 模块导入失败", str(e))
        return

    # 2. 单元测试：数字矛盾检测
    check(
        _has_number_contradiction("公差值 0.05", "公差值 0.1"),
        "数字 0.05 vs 0.1 → 矛盾",
    )
    check(
        not _has_number_contradiction("公差值 0.05", "公差值 0.05"),
        "数字相同 → 不矛盾",
    )
    check(
        not _has_number_contradiction("无数字文本", "另一段无数字文本"),
        "双方都无数字 → 不矛盾",
    )

    # 3. 单元测试：严格度检测（A 弱 B 强）
    check(
        _has_strict_vs_weaker("宜采用 Ra 3.2", "必须采用 Ra 1.6"),
        "A 弱（宜） B 强（必须）→ 增强",
    )
    check(
        not _has_strict_vs_weaker("必须 Ra 1.6", "宜 Ra 3.2"),
        "A 强 B 弱 → 非增强方向",
    )

    # 4. 单元测试：Jaccard 相似度
    s1 = _tokenize_zh("圆度公差带 半径差 两同心圆")
    s2 = _tokenize_zh("圆度公差带 半径差 两同心圆")
    check(_jaccard(s1, s2) == 1.0, "完全相同 → Jaccard=1.0")
    s3 = _tokenize_zh("完全不同的文本")
    check(_jaccard(s1, s3) < 0.2, "完全不同 → Jaccard<0.2")

    # 5. 端到端：构造 2 条矛盾条文
    from app.schemas.kb import ClauseRecord

    national_clauses = [
        ClauseRecord(
            standard="GB/T 1182-2018",
            clause_id="5.2",
            title="圆度公差",
            category="shape_tolerance",
            keywords=["圆度", "公差"],
            references=[],
            version="2018",
            is_sample=True,
            original_text="圆度公差带是在同一正截面上，半径差为公差值 0.1 的两同心圆之间的区域。",
            source_file="gb.md",
        ),
        ClauseRecord(
            standard="GB/T 1182-2018",
            clause_id="6.1",
            title="表面粗糙度",
            category="surface_parameter",
            keywords=["Ra", "粗糙度"],
            references=[],
            version="2018",
            is_sample=True,
            original_text="表面粗糙度宜采用 Ra 3.2。",
            source_file="gb.md",
        ),
        # 第三条：国标独有，企业无 → missing
        ClauseRecord(
            standard="GB/T 1182-2018",
            clause_id="7.1",
            title="圆柱度公差",
            category="shape_tolerance",
            keywords=["圆柱度", "公差"],
            references=[],
            version="2018",
            is_sample=True,
            original_text="圆柱度公差带是半径差为公差值 0.05 的两同轴圆柱面之间的区域。",
            source_file="gb.md",
        ),
    ]
    enterprise_clauses = [
        ClauseRecord(
            standard="Q/ENT 001-2024",
            clause_id="5.2",
            title="圆度公差",
            category="shape_tolerance",
            keywords=["圆度", "公差"],
            references=[],
            version="2024",
            is_sample=False,
            original_text="圆度公差带是在同一正截面上，半径差为公差值 0.05 的两同心圆之间的区域。",
            source_file="ent.docx",
        ),
        ClauseRecord(
            standard="Q/ENT 001-2024",
            clause_id="6.1",
            title="表面粗糙度",
            category="surface_parameter",
            keywords=["Ra", "粗糙度"],
            references=[],
            version="2024",
            is_sample=False,
            original_text="表面粗糙度必须采用 Ra 1.6。",
            source_file="ent.docx",
        ),
    ]

    # 6. 不使用 LLM 检测（关键词匹配）
    try:
        report = detect_conflicts(
            clauses_a=national_clauses,
            clauses_b=enterprise_clauses,
            standard_a="GB/T 1182-2018",
            standard_b="Q/ENT 001-2024",
            use_llm=False,
        )
        check(True, "detect_conflicts(use_llm=False) 执行成功")
        check(
            report.llm_used is False,
            "use_llm=False 时 llm_used=False",
        )
        check(
            report.total > 0,
            f"检测到 >= 1 条冲突（实际 {report.total}）",
            f"by_type={report.by_type}",
        )

        # 至少应检测到一条 contradiction（5.2 条款 0.1 vs 0.05）
        contradictions = [
            c for c in report.conflicts if c.conflict_type == "contradiction"
        ]
        check(
            len(contradictions) >= 1,
            f"检测到 >= 1 条 contradiction（实际 {len(contradictions)}）",
            f"items={[c.clause_a_id for c in contradictions]}",
        )

        # 至少应检测到一条 enhancement（6.1 条款 宜 vs 必须）
        enhancements = [
            c for c in report.conflicts if c.conflict_type == "enhancement"
        ]
        check(
            len(enhancements) >= 1,
            f"检测到 >= 1 条 enhancement（实际 {len(enhancements)}）",
            f"items={[c.clause_a_id for c in enhancements]}",
        )

        # 应检测到 missing（7.1 圆柱度 企业无）
        missing = [c for c in report.conflicts if c.conflict_type == "missing"]
        check(
            len(missing) >= 1,
            f"检测到 >= 1 条 missing（实际 {len(missing)}）",
            f"items={[c.clause_a_id for c in missing]}",
        )

        # 验证冲突项字段完整性
        if report.conflicts:
            c0 = report.conflicts[0]
            check(
                all([c0.standard_a, c0.standard_b, c0.conflict_type]),
                "冲突项字段完整（standard_a/standard_b/conflict_type）",
            )
            check(
                c0.detection_method in ("llm", "keyword", "both"),
                f"detection_method 合法（{c0.detection_method}）",
            )
    except Exception as e:
        check(False, "关键词检测执行异常", str(e))

    # 7. LLM 调用路径（LLM 不可用时应降级，不报错）
    try:
        report_llm = detect_conflicts(
            clauses_a=national_clauses,
            clauses_b=enterprise_clauses,
            standard_a="GB/T 1182-2018",
            standard_b="Q/ENT 001-2024",
            use_llm=True,
        )
        check(True, "detect_conflicts(use_llm=True) 执行成功（LLM 不可用时降级）")
        # LLM 不可用时 llm_used 仍为 False
        if report_llm.llm_used:
            _info("LLM 服务可用，本次使用了 LLM 检测")
        else:
            env_limit("LLM 服务不可用，已降级为纯关键词匹配（仍能检测到冲突）")
        check(
            report_llm.total > 0,
            f"LLM 路径检测到 >= 1 条冲突（实际 {report_llm.total}）",
        )
    except Exception as e:
        check(False, "LLM 路径执行异常", str(e))


# ---------------------------------------------------------------------------
# SubTask 14.3：多套规范配置切换
# ---------------------------------------------------------------------------


def test_standard_profiles() -> None:
    section("SubTask 14.3: 多套规范配置切换")

    # 1. 模块导入
    try:
        from app.services.kb.standard_profile import (
            StandardProfileManager,
            get_manager,
            DEFAULT_PROFILE_NAME,
            ENV_ACTIVE_PROFILE,
        )
        check(True, "standard_profile 模块导入成功")
    except Exception as e:
        check(False, "standard_profile 模块导入失败", str(e))
        return

    # 2. 后端初始化（PostgreSQL 不可用时降级 JSON）
    with tempfile.TemporaryDirectory(prefix="task14_profiles_") as tmpdir:
        json_path = Path(tmpdir) / "profiles.json"
        # 强制重置单例 + 指定 JSON 路径
        StandardProfileManager._instance = None  # type: ignore[attr-defined]
        try:
            mgr = StandardProfileManager(json_path=json_path)
        except Exception as e:
            check(False, "StandardProfileManager 初始化失败", str(e))
            return

        _info(f"后端：{mgr.backend_name}")
        check(
            mgr.backend_name in ("postgres", "json"),
            f"后端名称合法（{mgr.backend_name}）",
        )
        if mgr.backend_name == "json":
            _info("PostgreSQL 不可用，已降级为 JSON 文件后端")

        # 3. 默认配置自动种子
        profiles = mgr.list_profiles()
        check(
            len(profiles) >= 1,
            f"初始化后至少 1 个默认配置（实际 {len(profiles)}）",
        )
        check(
            any(p.name == DEFAULT_PROFILE_NAME for p in profiles),
            f"含 default 配置",
        )

        # 4. 默认配置应为活跃
        active = mgr.get_active_profile()
        check(
            active is not None and active.name == DEFAULT_PROFILE_NAME,
            f"默认活跃配置为 default（实际 {active.name if active else None}）",
        )

        # 5. 创建 2 个新配置
        p1 = mgr.create_profile(
            name="enterprise_a",
            standards=["Q/XX 001-2024", "GB/T 1182-2018"],
            description="企业 A 内控规范集",
            priority=20,
        )
        check(p1.name == "enterprise_a", "创建 enterprise_a 配置")
        check(
            "Q/XX 001-2024" in p1.standards,
            "enterprise_a 含 Q/XX 001-2024",
        )

        p2 = mgr.create_profile(
            name="industry_jb",
            standards=["JB/T 5000-2020", "GB/T 131-2006"],
            description="行业标准集",
            priority=15,
        )
        check(p2.name == "industry_jb", "创建 industry_jb 配置")

        # 6. 列表应含 3 个配置
        profiles = mgr.list_profiles()
        check(
            len(profiles) >= 3,
            f"列表含 >= 3 个配置（实际 {len(profiles)}）",
        )
        # 按 priority 降序排列：enterprise_a(20) > industry_jb(15) > default(10)
        check(
            profiles[0].name == "enterprise_a",
            f"按 priority 降序，首个为 enterprise_a（实际 {profiles[0].name}）",
        )

        # 7. 切换活跃配置
        ok = mgr.set_active_profile("enterprise_a")
        check(ok, "切换活跃配置到 enterprise_a 成功")
        active = mgr.get_active_profile()
        check(
            active is not None and active.name == "enterprise_a",
            f"当前活跃配置为 enterprise_a（实际 {active.name if active else None}）",
        )

        # 8. 验证 list 中 is_active 标记
        profiles = mgr.list_profiles()
        ea = next((p for p in profiles if p.name == "enterprise_a"), None)
        check(
            ea is not None and ea.is_active is True,
            "list_profiles 中 enterprise_a.is_active=True",
        )
        default = next((p for p in profiles if p.name == DEFAULT_PROFILE_NAME), None)
        check(
            default is not None and default.is_active is False,
            "list_profiles 中 default.is_active=False",
        )

        # 9. 环境变量覆盖活跃配置
        os.environ[ENV_ACTIVE_PROFILE] = "industry_jb"
        try:
            active = mgr.get_active_profile()
            check(
                active is not None and active.name == "industry_jb",
                f"环境变量 STANDARD_PROFILE=industry_jb 覆盖活跃配置（实际 {active.name if active else None}）",
            )
        finally:
            del os.environ[ENV_ACTIVE_PROFILE]

        # 10. 切换到不存在的配置应失败
        ok = mgr.set_active_profile("nonexistent")
        check(ok is False, "切换到不存在的配置返回 False")

        # 11. 删除配置
        ok = mgr.delete_profile("enterprise_a")
        check(ok, "删除 enterprise_a 配置成功")
        profiles = mgr.list_profiles()
        check(
            not any(p.name == "enterprise_a" for p in profiles),
            "删除后列表中无 enterprise_a",
        )

        # 12. 删除 default 应被拒绝
        ok = mgr.delete_profile(DEFAULT_PROFILE_NAME)
        check(ok is False, "删除 default 配置被拒绝")

    # 13. 清理单例
    StandardProfileManager._instance = None  # type: ignore[attr-defined]


# ---------------------------------------------------------------------------
# SubTask：API 端点导入
# ---------------------------------------------------------------------------


def test_api_endpoints() -> None:
    section("API 端点：路由注册与 schema")

    # 1. 导入 kb 端点模块（验证无语法/导入错误）
    try:
        from app.api.v1.endpoints import kb as kb_endpoint
        check(True, "kb 端点模块导入成功")
    except Exception as e:
        check(False, "kb 端点模块导入失败", str(e))
        return

    # 2. 验证新增端点已注册
    # 注意：路由列表可能含无 path 属性的对象（如 Mount），用 getattr 安全访问。
    routes: dict[str, Any] = {}
    for r in kb_endpoint.router.routes:
        rpath = getattr(r, "path", None)
        if rpath:
            routes[rpath] = getattr(r, "methods", set())
    expected_new_routes = [
        "/enterprise-standards/import",
        "/standards/conflicts",
        "/profiles",
        "/profiles/active",
    ]
    for path in expected_new_routes:
        # 路由前缀 /kb 由 router.py 挂载，端点模块内 path 不含前缀
        check(
            any(path in r for r in routes.keys()),
            f"端点已注册：{path}",
            f"actual_routes={list(routes.keys())}",
        )

    # 3. 验证 schema 可序列化
    try:
        from app.schemas.kb import (
            ConflictItem,
            ConflictReport,
            ConflictType,
            ConflictSeverity,
            EnterpriseImportResponse,
            StandardProfile,
            ProfileListResponse,
            ProfileCreateRequest,
            ProfileSetActiveRequest,
        )
        check(True, "Task 14 schema 全部导入成功")

        # 实例化测试
        item = ConflictItem(
            conflict_type="contradiction",
            severity="major",
            standard_a="GB/T 1182-2018",
            standard_b="Q/ENT 001-2024",
            clause_a_id="5.2",
            clause_b_id="5.2",
            description="测试冲突",
        )
        check(item.conflict_type == "contradiction", "ConflictItem 实例化")

        profile = StandardProfile(
            name="test",
            standards=["GB/T 1182-2018"],
            priority=10,
        )
        check(profile.name == "test", "StandardProfile 实例化")

        report = ConflictReport(
            standard_a="A",
            standard_b="B",
            conflicts=[item],
            total=1,
        )
        check(report.total == 1, "ConflictReport 实例化")

    except Exception as e:
        check(False, "schema 实例化失败", str(e))

    # 4. 验证完整 FastAPI app 可加载（无路由冲突）
    try:
        from app.main import app

        # FastAPI 0.100+ 使用 _IncludedRouter 包装嵌套路由，
        # app.routes 不再直接展开所有子路由的 path。
        # 改用 app.openapi() 获取所有已注册路径（权威来源）。
        schema = app.openapi()
        paths = set(schema.get("paths", {}).keys())
        kb_paths = sorted(p for p in paths if "/kb/" in p)
        check(
            "/api/v1/kb/enterprise-standards/import" in paths,
            "FastAPI app 含 /api/v1/kb/enterprise-standards/import 路由",
            f"kb_paths={kb_paths[:10]}",
        )
        check(
            "/api/v1/kb/standards/conflicts" in paths,
            "FastAPI app 含 /api/v1/kb/standards/conflicts 路由",
        )
        check(
            "/api/v1/kb/profiles" in paths,
            "FastAPI app 含 /api/v1/kb/profiles 路由",
        )
        check(
            "/api/v1/kb/profiles/active" in paths,
            "FastAPI app 含 /api/v1/kb/profiles/active 路由",
        )
    except Exception as e:
        check(False, "FastAPI app 加载失败", str(e))


# ===== 主入口 =====


def main() -> int:
    section("Task 14 企业规范自定义 - 自测开始")

    test_enterprise_import()
    test_conflict_detection()
    test_standard_profiles()
    test_api_endpoints()

    section("Task 14 自测汇总")
    total = _passed + _failed
    print(f"  通过: {_passed}/{total}", flush=True)
    print(f"  失败: {_failed}/{total}", flush=True)
    print(f"  环境限制: {_env_limits}", flush=True)
    if _failures:
        print("\n  失败明细:", flush=True)
        for f in _failures:
            print(f"    - {f}", flush=True)

    # 退出码：所有非环境限制的检查必须通过
    return 0 if _failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
